import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document(r'12. 엑셀의 정보를 불러와 수료증 자동생성\수료증양식.docx')

style = doc.styles['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph() 
run = para.add_run('\n\n') 
run = para.add_run('              제 2020-0001 호\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('수  료  증') 
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('        성       명: 장다인\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        생 년 월 일: 2017.12.12\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20) 
run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('2021.09.19') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('파이썬교육기관장') 
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'12. 엑셀의 정보를 불러와 수료증 자동생성\수료증결과.docx')
